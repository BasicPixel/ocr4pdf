import pytesseract
import sys
from pdf2image import convert_from_path
from docx import Document

# Use Homebrew to locate Tesseract automatically (if installed)
import os

if os.path.exists("/usr/local/Cellar/tesseract-ocr/"):
    tesseract_cmd = "/usr/local/Cellar/tesseract-ocr/latest/bin/tesseract"
else:
    # Provide a custom path if Tesseract isn't installed using Homebrew
    # Replace with your path
    pytesseract.pytesseract.tesseract_cmd = r'/opt/homebrew/bin/tesseract'

# Get PDF filename from command line arguments
if len(sys.argv) < 2:
    print("Error: Please provide the PDF filename as a command-line argument.")
    exit(1)
pdf_path = sys.argv[1]

# Dynamically create output filename
filename = os.path.splitext(pdf_path)[0]
output_filename = filename + '.txt'

# Convert PDF pages to images (one image per page)
images = convert_from_path(pdf_path, dpi=200)

# Create word document to store text
document = Document("template.docx")
document.add_heading(filename, 0)

# Extract text from each image and append it to the word document
for image in images:
    p = document.add_paragraph(
        pytesseract.image_to_string(image, lang='ara') + "\n")

    document.add_page_break()

document.save('./' + filename + '.docx')

print(f"Text extracted and saved to {output_filename}.")
